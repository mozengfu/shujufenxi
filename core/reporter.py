"""Word 报告生成模块"""
import tempfile
from pathlib import Path
from typing import Dict, Any
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from matplotlib.figure import Figure


class WordReporter:
    """Word 报告生成器"""

    def __init__(self):
        self.doc = None

    def create_document(self, title: str = '', author: str = '') -> Document:
        """创建新文档"""
        self.doc = Document()

        # 添加页眉页脚
        if title or author:
            self._add_header_footer(title, author)

        return self.doc

    def _add_header_footer(self, title: str = '', author: str = '') -> None:
        """添加页眉页脚"""
        # 页眉
        section = self.doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = title if title else '数据分析报告'
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header_para.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)

        # 页脚
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加页码
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # 页码域代码
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)

        run = footer_para.add_run('PAGE')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        run._r.append(instrText)

        run = footer_para.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)

        # 添加 / 总页数
        footer_para.add_run(' / ')

        run = footer_para.add_run()
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar3)

        run = footer_para.add_run('NUMPAGES')
        instrText2 = OxmlElement('w:instrText')
        instrText2.text = 'NUMPAGES'
        run._r.append(instrText2)

        run = footer_para.add_run()
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar4)

        for run in footer_para.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)

    def add_title(self, title: str, level: int = 0) -> None:
        """添加标题"""
        if self.doc is None:
            self.create_document()

        if level == 0:
            heading = self.doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            self.doc.add_heading(title, level=level)

    def add_paragraph(self, text: str) -> None:
        """添加段落"""
        if self.doc is None:
            self.create_document()
        self.doc.add_paragraph(text)

    def _set_cell_shading(self, cell, color: str) -> None:
        """设置单元格背景色"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._element.get_or_add_tcPr().append(shading_elm)

    def _set_cell_border(self, cell) -> None:
        """设置单元格边框"""
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)
        tcPr.append(tcBorders)

    def _set_cell_font(self, cell, bold: bool = False, size: int = 10, color: str = '000000') -> None:
        """设置单元格字体"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = bold
                run.font.size = Pt(size)
                run.font.color.rgb = RGBColor.from_string(color)

    def add_table_from_df(self, df: pd.DataFrame, header: bool = True, 
                          striped: bool = True, auto_width: bool = True) -> None:
        """从 DataFrame 添加表格（带样式美化）"""
        if self.doc is None:
            self.create_document()

        rows = len(df) + (1 if header else 0)
        cols = len(df.columns)

        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        if header:
            for idx, col_name in enumerate(df.columns):
                cell = table.rows[0].cells[idx]
                cell.text = str(col_name)
                # 表头样式：蓝色背景，白色字体，居中
                self._set_cell_shading(cell, '366092')
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(255, 255, 255)
                self._set_cell_border(cell)

        # 数据
        for row_idx, (_, row) in enumerate(df.iterrows(), 1 if header else 0):
            # 隔行变色
            if striped and row_idx > 0:
                row_color = 'F2F2F2' if (row_idx - 1) % 2 == 1 else 'FFFFFF'
            else:
                row_color = 'FFFFFF'

            for col_idx, value in enumerate(row):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = str(value) if pd.notna(value) else ''

                # 应用背景色
                if row_idx > 0 or not header:
                    self._set_cell_shading(cell, row_color)

                # 对齐方式
                for paragraph in cell.paragraphs:
                    # 数字右对齐，其他居中
                    if isinstance(value, (int, float)):
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # 字体大小
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

                self._set_cell_border(cell)

    def add_stats_report(self, stats: pd.DataFrame, title: str = '描述性统计报告') -> None:
        """添加统计报告"""
        if self.doc is None:
            self.create_document()

        self.add_title(title, level=1)
        self.add_paragraph('')
        self.add_table_from_df(stats)
        self.add_paragraph('')

    def add_quality_report(self, quality_report: Dict[str, Any]) -> None:
        """添加数据质量报告"""
        if self.doc is None:
            self.create_document()

        self.add_title('数据质量报告', level=1)
        self.add_paragraph('')

        # 缺失值
        if quality_report.get('missing'):
            self.add_title('缺失值检测', level=2)
            missing_data = []
            for col, info in quality_report['missing'].items():
                missing_data.append({
                    '列名': col,
                    '缺失数量': info['count'],
                    '缺失百分比': f"{info['percentage']}%"
                })
            if missing_data:
                self.add_table_from_df(pd.DataFrame(missing_data))
            self.add_paragraph('')

        # 异常值
        if quality_report.get('outliers'):
            self.add_title('异常值检测', level=2)
            outlier_data = []
            for col, info in quality_report['outliers'].items():
                outlier_data.append({
                    '列名': col,
                    '异常数量': info['count'],
                    '下界': info.get('lower_bound'),
                    '上界': info.get('upper_bound')
                })
            if outlier_data:
                self.add_table_from_df(pd.DataFrame(outlier_data))
            self.add_paragraph('')

        # 重复行
        if quality_report.get('duplicates'):
            dup = quality_report['duplicates']
            self.add_title('重复行检测', level=2)
            self.add_paragraph(f"重复行数: {dup['count']} ({dup['percentage']}%)")
            if dup.get('rows'):
                self.add_paragraph(f"行号: {', '.join(map(str, dup['rows'][:50]))}")
            self.add_paragraph('')

    def add_comparison_report(self, comparison: Dict[str, Any]) -> None:
        """添加对比分析报告"""
        if self.doc is None:
            self.create_document()

        self.add_title('对比分析报告', level=1)
        self.add_paragraph('')

        if 'key_match_count' in comparison:
            self.add_paragraph(f"关键列匹配数: {comparison['key_match_count']}")
            self.add_paragraph(f"仅在表1: {comparison.get('only_in_df1_rows', 0)} 行")
            self.add_paragraph(f"仅在表2: {comparison.get('only_in_df2_rows', 0)} 行")
            self.add_paragraph(f"两表都有: {comparison.get('in_both_rows', 0)} 行")
        else:
            self.add_paragraph(f"行数差异: {comparison.get('row_diff', 0)}")

        self.add_paragraph('')

        # 列差异
        if comparison.get('only_in_df1') or comparison.get('only_in_df2'):
            self.add_title('列差异', level=2)
            if comparison.get('only_in_df1'):
                self.add_paragraph(f"仅在表1: {', '.join(comparison['only_in_df1'])}")
            if comparison.get('only_in_df2'):
                self.add_paragraph(f"仅在表2: {', '.join(comparison['only_in_df2'])}")
            self.add_paragraph('')

    def add_chart(self, fig: Figure, title: str = '图表') -> None:
        """将 matplotlib 图表嵌入文档"""
        if self.doc is None:
            self.create_document()

        # 先用临时文件保存图片，再插入
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            tmp_path = Path(tmp.name)
        try:
            fig.savefig(tmp_path, dpi=150, bbox_inches='tight')
            self.add_title(title, level=1)
            self.doc.add_picture(str(tmp_path), width=Inches(5.5))
            self.add_paragraph('')
        finally:
            if tmp_path.exists():
                tmp_path.unlink()

    def save(self, file_path: str) -> None:
        """保存文档"""
        if self.doc is None:
            raise ValueError("文档未创建")

        path = Path(file_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(file_path)